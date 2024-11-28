import os
from docx import Document
from datetime import datetime
from autogen import ConversableAgent

from dotenv import load_dotenv

load_dotenv()

model = "gpt-3.5-turbo"
llm_config = {
    "model": model,
    "temperature": 0.0,
    "api_key": os.environ["OPENAI_API_KEY"],
}

# Load the Word template
def load_template(filepath):
    return Document(filepath)

# Fill in the template including paragraph and table
def fill_template(doc, fields):
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in fields.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in fields.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", value)
    return doc

# Save the filled report
def save_report(doc, output_filepath):
    doc.save(output_filepath)

# Example function to simulate data for a table-based template
############## need to do works here to fetch the data ##############
def fetch_report_data_with_table():
    return {
        "observation": "Observation 1: Copy observation verbatim, including annotation, if any.",
        "response": "Response: Corrective actions will be completed by December 1, 2024.",
        "completed_actions": """
        On November 10, 2024, COMPANY completed corrective actions. See Appendix 2, Attachment A for a copy of the revised records.
        On November 15, 2024, COMPANY updated training procedures. See Appendix 2, Attachment B for the new training log.
        """,
        "planned_actions": """
        By December 15, 2024, COMPANY will implement new policies.
        By December 30, 2024, COMPANY will conduct a follow-up audit to ensure compliance.
        COMPANY considers this item to be closed.
        """,
    }


# Define Conversable Agent
agent = ConversableAgent(
    name="ReportAgent",
    system_message="You are an agent generating a report.",
    llm_config=llm_config
)

# Generate the report using the updated template with tables
def generate_report_with_table_agent(template_path, output_path):
    # Step 1: Load the Word template
    doc = load_template(template_path)
    
    # Step 2: Fetch report data
    data = fetch_report_data_with_table()
    
    # Step 3: Simulate agent processing or refining specific fields (optional)
    prompt = f"""
    Please review the following data and provide a concise summary for 'planned_actions':
    {data['planned_actions']}
    """
    response = agent.generate_reply(messages=[{"role": "user", "content": prompt}])
    # Debugging: Check the structure of the response
    print("Response type:", type(response))
    print("Response content:", response)
    refined_planned_actions = response["content"] if isinstance(response, dict) else response
    
    # Step 4: Update the 'planned_actions' field with the agent's response
    data["planned_actions"] = refined_planned_actions
    
    # Step 5: Fill in the Word template with data
    filled_doc = fill_template(doc, data)
    
    # Step 6: Save the generated report
    save_report(filled_doc, output_path)
    print(f"Report generated and saved to: {output_path}")



if __name__ == "__main__":
    template_path = "./documents/fda_template_with_table.docx"  # Path to your Word template
    output_path = "./documents/fda_generated_report_with_table.docx"  # Path to save the generated report
    
    generate_report_with_table_agent(template_path, output_path)
