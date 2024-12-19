from pathlib import Path
import streamlit as st
import logging
from agents.agent_manager import group_chat
from agents.conversation_workflow import conversation_workflow
import asyncio
from docx import Document
import io
import os
import shutil
import stat

# Initialize logging
# logging.basicConfig(level=logging.INFO)

IMAGE_DIR = Path(__file__).parent.resolve() / "images"
assert IMAGE_DIR.is_dir()
berkeley_logo_path = IMAGE_DIR / "berkeley_logo.png"
assert berkeley_logo_path.is_file()
berkeley_hackathon_banner_path = IMAGE_DIR / "berkeley_hackathon_banner.png"
assert berkeley_hackathon_banner_path.is_file()


# Function to check allowed file extensions
def allowed_file(filename):
    ALLOWED_EXTENSIONS = {"txt", "pdf", "docx"}
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# Function to read file contents
def read_file_contents(file):
    filename = file.name.lower()
    try:
        if filename.endswith(".txt"):
            content = file.read()
            return content.decode("utf-8") if isinstance(content, bytes) else content
        elif filename.endswith(".pdf"):
            return extract_text_from_pdf(file)
        elif filename.endswith(".docx"):
            return extract_text_from_docx(file)
        else:
            return None
    except Exception as e:
        st.error(f"Error reading {filename}: {e}")
        return None


# Function to extract text from PDF files
def extract_text_from_pdf(file):
    import PyPDF2

    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return None


# Function to extract text from DOCX files
def extract_text_from_docx(file):
    from docx import Document
    from io import BytesIO

    try:
        file.seek(0)
        document = Document(BytesIO(file.read()))
        text = "\n".join([para.text for para in document.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return None


# Function to run the workflow synchronously
def run_workflow(group_chat):
    # Run the conversation workflow
    try:
        with st.spinner("🤖 AI is processing your request..."):
            result = asyncio.run(conversation_workflow(group_chat))
        return result
    except Exception as e:
        st.error(f"An error occurred during processing: {e}")
        # logging.error(f"Error during workflow: {e}")
        return None


def create_word_file(content):
    document = Document()
    document.add_paragraph(content)
    return document


def get_base64_image(image_path):
    import base64

    with open(image_path, "rb") as f:
        data = f.read()
    return base64.b64encode(data).decode()


def main():

    # Remove .cache file or directory
    cache_path = ".cache"
    if os.path.exists(cache_path):
        try:
            # Change the permissions of the file/directory to ensure it can be removed
            os.chmod(cache_path, stat.S_IRWXU | stat.S_IRWXG | stat.S_IRWXO)

            if os.path.isfile(cache_path):
                os.remove(cache_path)
                logging.info(f"Successfully removed file {cache_path}")
            elif os.path.isdir(cache_path):
                shutil.rmtree(cache_path)
                logging.info(f"Successfully removed directory {cache_path}")
            else:
                logging.warning(f"{cache_path} is neither a file nor a directory.")
        except Exception as e:
            logging.error(f"Error removing {cache_path}: {e}")
    else:
        logging.info(f"No cache file or directory found at {cache_path}")
    # Set page configuration
    st.set_page_config(
        page_title="FDA Warning Letter Processor",
        page_icon="📄",
        layout="centered",
        initial_sidebar_state="expanded",
    )

    # Custom CSS for Berkeley colors and black font
    berkeley_blue = "#003262"
    california_gold = "#FDB515"

    st.markdown(
        f"""
        <style>
        .stApp {{
            background-color: white;
        }}
        h1, h2, h3, h4 {{
            color: {berkeley_blue};
        }}
        body, p, div {{
            color: black;
        }}
        .stButton>button {{
            background-color: {california_gold};
            color: black;
            border: none;
        }}
        /* Ensure image captions are black */
        .stImage figcaption {{
            color: black;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Add the Berkeley Hackathon banner at the top
    st.image(str(berkeley_hackathon_banner_path), use_container_width=True)

    st.title("📄 FDA Warning Letter Processor")
    st.markdown(
        "Welcome to **RegulAItor**! Please upload the warning letter and the template files to get started."
    )

    # Add Hackathon and team information
    st.header("🚀 Berkeley AI Hackathon 2024 Submission")
    st.markdown(
        """
    Join us at the **world's largest AI Hackathon** hosted by **Berkeley** Large Language Model Agents **MOOC, Fall 2024**! We're excited to present our project and contribute to innovative solutions during this event.
    """
    )

    # Display team members
    st.header("👥 Meet the Team")
    team_members = {
        "Mobin AZIMIPANAH": IMAGE_DIR / "mobin.jpeg",
        "Joan Cheung, PharmD": IMAGE_DIR / "joan.jpeg",
    }
    cols = st.columns(len(team_members))

    for idx, (name, photo) in enumerate(team_members.items()):
        with cols[idx]:
            st.markdown(
                "<div style='background-color: #f0f0f0; padding: 10px; text-align: center;'>",
                unsafe_allow_html=True,
            )
            st.image(photo, width=150)
            st.markdown(
                f"<p style='color:black; font-weight:bold;'>{name}</p>",
                unsafe_allow_html=True,
            )
            st.markdown("</div>", unsafe_allow_html=True)

    # Add a sidebar with some information
    st.sidebar.title("ℹ️ Instructions")
    st.sidebar.info(
        """
        **How to use this app:**

        1. **Upload** the FDA warning letter (TXT, PDF, DOCX).
        2. **Upload** your corrective action plan template (TXT, PDF, DOCX).
        3. Click on **Start Processing** to generate the corrective action plan.
        4. **Download** and review the generated Corrective Action Plan.
        """
    )

    st.sidebar.title("ℹ️ About")
    st.sidebar.info(
        """
        **RegulAItor** helps you process FDA warning letters and generate corrective action plans using AI. Streamline your compliance workflow and ensure accuracy with advanced AI capabilities.

        **Team Innovators**
        Submitted for the UC Berkeley AI Hackathon 2024.
        """
    )

    # Sidebar Hackathon logo
    st.sidebar.image(berkeley_logo_path, use_container_width=True)

    # Initialize session state variables
    if "stage" not in st.session_state:
        st.session_state.stage = "upload"
    if "warning_letter_content" not in st.session_state:
        st.session_state.warning_letter_content = None
    if "template_content" not in st.session_state:
        st.session_state.template_content = None

    # Stage: Upload
    if st.session_state.stage == "upload":
        st.header("Step 1: Upload Files")
        col1, col2 = st.columns(2)
        with col1:
            warning_letter_file = st.file_uploader(
                "📄 Upload FDA Warning Letter",
                type=["txt", "pdf", "docx"],
                key="warning_letter",
            )
        with col2:
            template_file = st.file_uploader(
                "📄 Upload Corrective Action Plan Template",
                type=["txt", "pdf", "docx"],
                key="template",
            )

        if warning_letter_file and template_file:
            # Validate and read files
            if not allowed_file(warning_letter_file.name):
                st.error(
                    "⚠️ Invalid warning letter file type. Only TXT, PDF, and DOCX files are allowed."
                )
                return
            if not allowed_file(template_file.name):
                st.error(
                    "⚠️ Invalid template file type. Only TXT, PDF, and DOCX files are allowed."
                )
                return

            st.session_state.warning_letter_content = read_file_contents(
                warning_letter_file
            )
            st.session_state.template_content = read_file_contents(template_file)

            if (
                st.session_state.warning_letter_content
                and st.session_state.template_content
            ):
                st.header("Step 2: Process Warning Letter")
                process_button = st.button("🚀 Start Processing")

                if process_button:
                    # Set context for the group chat
                    group_chat.context = {
                        "warning_letter": st.session_state.warning_letter_content,
                        "template": st.session_state.template_content,
                    }

                    # Run the workflow synchronously and provide updates
                    result = run_workflow(group_chat)

                    if result.get(True):
                        st.balloons()
                        st.success("✅ Processing completed successfully!")
                        approved_corrective_action_plan = group_chat.context.get(
                            "Approved_corrective_action_plan"
                        )
                        if approved_corrective_action_plan:
                            st.header("Step 3: Download Corrective Action Plan")
                            word_file = create_word_file(
                                approved_corrective_action_plan
                            )
                            with io.BytesIO() as output:
                                word_file.save(output)
                                data = output.getvalue()
                            st.download_button(
                                label="💾 Download Corrective Action Plan",
                                data=data,
                                file_name="corrective_action_plan.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                            with st.expander("👀 Preview Corrective Action Plan"):
                                st.markdown(
                                    approved_corrective_action_plan.replace(
                                        "\n", "  \n"
                                    ),
                                    unsafe_allow_html=True,
                                )
                    elif result.get(False):
                        st.error(
                            "⚠️ User input is not valid. Please review the warning letter or template."
                        )
                    else:
                        st.error("⚠️ An error occurred during processing.")
            else:
                st.error("⚠️ Could not read the uploaded files. Please try again.")

    # Add a footer with Hackathon and Berkeley logo
    st.markdown(
        """
    <hr>
    <div style='text-align: center;'>
        <img src='data:image/png;base64,{}' width='150'>
        <p>Submitted for the UC Berkeley AI Hackathon 2024</p>
        <p>Developed with ❤️ by Team Innovators</p>
    </div>
    """.format(
            get_base64_image(str(berkeley_logo_path))
        ),
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()
